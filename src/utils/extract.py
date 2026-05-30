"""Document text extraction endpoint — turns an uploaded PDF/DOCX/TXT/image into plain text."""

import base64
import io

import httpx
from fastapi import APIRouter, File, HTTPException, UploadFile

from src.utils.chat import _api_key, _azure_endpoint, _model

router = APIRouter(prefix="/api")

MAX_BYTES = 10 * 1024 * 1024  # 10 MB cap

IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".webp", ".gif")

VISION_PROMPT = (
    "Extract ALL text from this image verbatim. Preserve structure — headings, "
    "lists, and tables. If the image is a diagram or flowchart, also describe the "
    "boxes and how they connect, in plain text. Output only the extracted content, "
    "no commentary or preamble."
)


async def _extract_image(data: bytes, content_type: str) -> str:
    """Use the vision-capable LLM to transcribe text from an image."""
    b64 = base64.b64encode(data).decode()
    data_url = f"data:{content_type or 'image/png'};base64,{b64}"

    async with httpx.AsyncClient(timeout=90) as client:
        res = await client.post(
            f"{_azure_endpoint()}/chat/completions",
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {_api_key()}",
            },
            json={
                "model": _model(),
                "messages": [
                    {"role": "system", "content": VISION_PROMPT},
                    {"role": "user", "content": [
                        {"type": "text", "text": "Extract the text from this image."},
                        {"type": "image_url", "image_url": {"url": data_url}},
                    ]},
                ],
            },
        )
        res.raise_for_status()
        return (res.json()["choices"][0]["message"]["content"] or "").strip()


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # include table cell text too
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore").strip()


@router.post("/extract")
async def extract(file: UploadFile = File(...)):
    data = await file.read()
    if len(data) > MAX_BYTES:
        raise HTTPException(status_code=413, detail="File too large (max 10 MB)")

    name = (file.filename or "").lower()

    try:
        if name.endswith(".pdf"):
            text = _extract_pdf(data)
        elif name.endswith(".docx"):
            text = _extract_docx(data)
        elif name.endswith((".txt", ".md")):
            text = _extract_txt(data)
        elif name.endswith(IMAGE_EXTS):
            text = await _extract_image(data, file.content_type)
        else:
            raise HTTPException(
                status_code=415,
                detail="Unsupported file type. Upload PDF, DOCX, TXT, MD, or an image.",
            )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Could not read file: {exc}")

    if not text:
        raise HTTPException(
            status_code=422,
            detail="No readable text found. If this is a scanned PDF, it needs OCR.",
        )

    return {"filename": file.filename, "text": text, "chars": len(text)}
